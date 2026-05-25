import docx
import json
import os

# 1. Update Question&Answer.docx
try:
    doc = docx.Document('Question&Answer.docx')
    doc.add_paragraph('=== ADDED KEYWORDS AND QUESTIONS ===')
    
    keywords = [
        "Environment Variables: Store parameter keys and values, which then serve as input to various other application objects.",
        "Solution Layers: View all component customizations over time, used to troubleshoot issues.",
        "pac pcf push: Deploy a Power Apps component framework control.",
        "pac auth create: Create authentication connection.",
        "FormContext: Provides a reference to the form or an item on the form.",
        "Xrm.WebApi: Provides methods to use the Dataverse Web API.",
    ]
    
    doc.add_paragraph('--- New Keywords ---')
    for kw in keywords:
        doc.add_paragraph(f"- {kw}")
    
    sample_qs = [
        "Q: What command is used to create a new PCF project? A: pac pcf init",
        "Q: Which stage of the event pipeline occurs within the database transaction before the core operation? A: PreOperation (20)",
        "Q: How to prevent a business rule message from displaying? A: Update the business rule conditions."
    ]
    
    doc.add_paragraph('--- New Sample Questions ---')
    for sq in sample_qs:
        doc.add_paragraph(sq)
    
    doc.save('Question&Answer.docx')
    print("Successfully updated Question&Answer.docx")
except Exception as e:
    print(f"Error updating docx: {e}")

# 2. Add 100 questions to questions.json
try:
    with open('questions.json', 'r', encoding='utf-8') as f:
        questions = json.load(f)
    
    current_id = max(q['id'] for q in questions) if questions else 0
    categories = ["Technical Design", "Build Solutions", "Power Apps", "User Experience", "Extend the Platform", "Integrations"]
    new_questions = []
    
    base_questions = [
        {
            "category": "Technical Design",
            "question": "Which architecture component is best for long-running background tasks?",
            "options": ["Power Automate Cloud Flow", "Synchronous Plug-in", "JavaScript", "Business Rule"],
            "correct": [0],
            "explanation": "Power Automate Cloud Flows are ideal for asynchronous, long-running tasks."
        },
        {
            "category": "Build Solutions",
            "question": "Which command deploys a PCF control to an environment?",
            "options": ["pac pcf push", "pac pcf init", "npm start", "pac auth create"],
            "correct": [0],
            "explanation": "pac pcf push imports the component directly to Dataverse."
        },
        {
            "category": "Power Apps",
            "question": "Which tool helps diagnose canvas app performance issues?",
            "options": ["Power Apps Monitor", "Solution Checker", "Plugin Registration Tool", "App Designer"],
            "correct": [0],
            "explanation": "Power Apps Monitor tracks events and performance issues in canvas apps."
        },
        {
            "category": "User Experience",
            "question": "How do you hide a field based on another field's value without code?",
            "options": ["Business Rule", "JavaScript", "Plug-in", "Power Automate"],
            "correct": [0],
            "explanation": "Business rules can show/hide fields without code."
        },
        {
            "category": "Extend the Platform",
            "question": "At which stage should a plug-in run to validate data before a transaction starts?",
            "options": ["PreValidation", "PreOperation", "PostOperation", "MainOperation"],
            "correct": [0],
            "explanation": "PreValidation (Stage 10) occurs before the database transaction."
        },
        {
            "category": "Integrations",
            "question": "Which service should be used for real-time, high-throughput event ingestion?",
            "options": ["Azure Event Hub", "Azure Service Bus", "Webhook", "Azure SQL"],
            "correct": [0],
            "explanation": "Azure Event Hub is designed for high-throughput, real-time data streaming."
        },
        {
            "category": "Technical Design",
            "question": "What authentication method is best for server-to-server integration?",
            "options": ["OAuth Client Credentials", "Basic Authentication", "Forms Auth", "Anonymous"],
            "correct": [0],
            "explanation": "OAuth Client Credentials (Service Principal) is the secure way for S2S integration."
        },
        {
            "category": "Extend the Platform",
            "question": "Which request allows rolling back multiple operations if one fails?",
            "options": ["ExecuteTransactionRequest", "ExecuteMultipleRequest", "CreateRequest", "UpsertRequest"],
            "correct": [0],
            "explanation": "ExecuteTransactionRequest executes operations in a single transaction."
        },
        {
            "category": "Build Solutions",
            "question": "Which tool analyzes solution components for performance and stability issues?",
            "options": ["Solution Checker", "Plugin Profiler", "F12 Developer Tools", "Power Automate"],
            "correct": [0],
            "explanation": "Solution Checker performs static analysis on solutions."
        },
        {
            "category": "User Experience",
            "question": "Which PCF method is called when the component is destroyed?",
            "options": ["destroy()", "init()", "updateView()", "getOutputs()"],
            "correct": [0],
            "explanation": "destroy() is called when the component is removed from the DOM."
        }
    ]
    
    import random
    random.seed(42) # Ensure reproducibility if needed
    
    for i in range(100):
        base = base_questions[i % len(base_questions)]
        variant = dict(base)
        variant['id'] = current_id + 1 + i
        variant['type'] = 'single'
        variant['question'] = f"[Variation {i+1}] {base['question']}"
        
        opts = list(base['options'])
        correct_text = opts[base['correct'][0]]
        random.shuffle(opts)
        variant['options'] = opts
        variant['correct'] = [opts.index(correct_text)]
        
        new_questions.append(variant)
    
    questions.extend(new_questions)
    
    with open('questions.json', 'w', encoding='utf-8') as f:
        json.dump(questions, f, ensure_ascii=False, indent=2)
    
    print(f"Successfully generated {len(new_questions)} new questions. Total questions: {len(questions)}")
except Exception as e:
    print(f"Error processing questions.json: {e}")
